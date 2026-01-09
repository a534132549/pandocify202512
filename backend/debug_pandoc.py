import pypandoc
import os
from docx import Document

# Markdown with "tight" lists (no newline before list)
test_md = """Start of paragraph.
* List Item 1
    * List Item 1.1
* List Item 2
"""

output_filename = "debug_test.docx"
reference_filename = "reference.docx"

if os.path.exists(output_filename):
    os.remove(output_filename)

print("Converting using default 'markdown'...")
try:
    pypandoc.convert_text(
        test_md,
        'docx',
        format='markdown+tex_math_dollars+hard_line_breaks',
        outputfile=output_filename,
        extra_args=['--reference-doc=' + reference_filename] if os.path.exists(reference_filename) else []
    )
    
    # Inspect the docx
    doc = Document(output_filename)
    print("--- Docx Paragraphs ---")
    for p in doc.paragraphs:
        # Check if style name contains 'List' or if text starts with '*'
        print(f"Style: '{p.style.name}', Text: '{p.text}'")
        
except Exception as e:
    print(f"Error: {e}")

print("\nConverting using 'gfm'...")
try:
    pypandoc.convert_text(
        test_md,
        'docx',
        format='gfm+tex_math_dollars', # hard_line_breaks might not be standard in gfm?
        outputfile="debug_test_gfm.docx",
        extra_args=['--reference-doc=' + reference_filename] if os.path.exists(reference_filename) else []
    )
    
    doc = Document("debug_test_gfm.docx")
    print("--- GFM Docx Paragraphs ---")
    for p in doc.paragraphs:
        print(f"Style: '{p.style.name}', Text: '{p.text}'")

except Exception as e:
    print(f"Error: {e}")
